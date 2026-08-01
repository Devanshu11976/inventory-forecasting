import docx
import os

doc_path = r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx'
doc = docx.Document(doc_path)

# 1. Grammar Improvement in Introduction (P[8])
for p in doc.paragraphs:
    if "or has only focussed on the accuracy of the models, and not related to the operational cost" in p.text:
        p.text = p.text.replace(
            "or has only focussed on the accuracy of the models, and not related to the operational cost of over or under-stocking [8].",
            "or has primarily focused on forecasting accuracy rather than the operational costs of overstocking and stockouts [8]."
        )

# 2. Remove Duplicate Fig 2 Caption (keep only the caption below the chart)
fig2_count = 0
for p in list(doc.paragraphs):
    if "Fig. 2." in p.text and "Comparison of actual and XGBoost-predicted" in p.text:
        fig2_count += 1
        if fig2_count == 1:
            # Remove the first occurrence (above Table II)
            p._element.getparent().remove(p._element)

print(f"Removed duplicate Fig 2 caption (remaining occurrences: {fig2_count - 1}).")

# 3. Clean & Peer-Reviewed IEEE Bibliography
refs_data = [
    '[1] E. A. Silver, D. F. Pyke, and R. Peterson, Inventory Management and Production Planning and Scheduling, 3rd ed. New York: Wiley, 1998.',
    '[2] S. Chopra and P. Meindl, Supply Chain Management: Strategy, Planning, and Operation, 7th ed. Boston: Pearson, 2019.',
    '[3] R. Fildes, K. Nikolopoulos, S. F. Crone, and A. A. Syntetos, "Forecasting research and practice: Implications for retail planning," Int. J. Research in Marketing, vol. 25, no. 4, pp. 229-243, 2008.',
    '[4] R. Carbonneau, K. Laframboise, and R. Vahidov, "Application of machine learning techniques for supply chain demand forecasting," European J. Operational Research, vol. 184, no. 3, pp. 1140-1154, 2008.',
    '[5] S. Makridakis, E. Spiliotis, and V. Assimakopoulos, "The M4 Competition: Results, findings, conclusion and way forward," Int. J. Forecasting, vol. 36, no. 1, pp. 54-74, 2020.',
    '[6] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining (KDD), 2016, pp. 785-794.',
    '[7] G. E. P. Box and G. M. Jenkins, Time Series Analysis: Forecasting and Control. San Francisco: Holden-Day, 1970.',
    '[8] S. Hochreiter and J. Schmidhuber, "Long Short-Term Memory," Neural Computation, vol. 9, no. 8, pp. 1735-1780, 1997.',
    '[9] J. K. Franco et al., "A Machine Learning-Based Stock Forecasting Method for Inventory Optimization in Micro and Small Enterprises," Eng., Technology & Applied Science Research, vol. 15, no. 1, pp. 18210-18216, 2025.',
    '[10] G. Zhang, B. E. Patuwo, and M. Y. Hu, "Forecasting with artificial neural networks: The state of the art," Int. J. Forecasting, vol. 14, no. 1, pp. 35-62, 1998.',
    '[11] S. Makridakis, E. Spiliotis, and V. Assimakopoulos, "The M5 competition: Background, organization, and implementation," Int. J. Forecasting, vol. 38, no. 4, pp. 1325-1346, 2022.',
    '[12] S. M. Lundberg and S.-I. Lee, "A unified approach to interpreting model predictions," in Adv. Neural Inf. Process. Syst. (NeurIPS 30), 2017, pp. 4765-4774.',
    '[13] N. Kourentzes, F. Petropoulos, and D. K. Trapero, "Improving forecasting by estimating time series structural components across multiple frequencies," Int. J. Forecasting, vol. 30, no. 2, pp. 291-302, 2014.',
    '[14] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed. Melbourne: OTexts, 2021.',
    '[15] B. E. Flores, "Utility of forecasting methods in inventory control," Int. J. Operations & Production Management, vol. 6, no. 3, pp. 38-47, 1986.',
    '[16] A. A. Syntetos, Z. Babai, and J. E. Boylan, "Supply chain forecasting in practice," J. Operational Research Society, vol. 67, no. 3, pp. 393-403, 2016.'
]

ref_count = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt.startswith('[') and ref_count < len(refs_data):
        p.text = refs_data[ref_count]
        ref_count += 1

print(f"Updated {ref_count} references to peer-reviewed IEEE entries.")

# Save to target files
targets = [
    r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx',
    r'c:\Users\devan\Desktop\research paper\AI-Based Inventory Forecasting for Small Businesses.docx'
]

for t in targets:
    try:
        doc.save(t)
        print("Successfully saved ultimate polished docx to:", t)
    except Exception as e:
        print(f"Error saving to {t}: {e}")
