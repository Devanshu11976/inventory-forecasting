import docx
import os

doc_path = r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx'
doc = docx.Document(doc_path)

# 1. Remove duplicate Fig. 1 caption (P[32] or P[34])
fig1_count = 0
for p in list(doc.paragraphs):
    if "Fig. 1." in p.text and "System architecture:" in p.text:
        fig1_count += 1
        if fig1_count > 1:
            # Delete duplicate paragraph
            p._element.getparent().remove(p._element)

print(f"Cleaned duplicate Fig 1 captions (remaining count: {fig1_count - 1 if fig1_count > 1 else 1}).")

# 2. Tone down / remove exact dollar-value estimates (P[75], P[76], P[77])
for p in doc.paragraphs:
    if "74,710" in p.text or "205,477" in p.text:
        p.text = (
            "To make this concrete, applying the 95% service level safety-stock formulation (z = 1.65, 5-day lead time) "
            "directly to the representative stores evaluated here demonstrates substantial reductions in required safety buffers: "
            "safety-stock levels drop from 6,500.9 to 2,765.4 units at Store 1 (a 57.5% reduction), "
            "from 13,869.1 to 3,595.3 units at Store 4 (a 74.1% reduction), "
            "and from 12,069.2 to 5,486.8 units at Store 9 (a 54.5% reduction)."
        )
    elif "123.9M" in p.text or "working capital released" in p.text:
        p.text = (
            "Extrapolating the average 62.0% safety-stock reduction across all 1,115 stores demonstrates that "
            "decreasing forecast error significantly reduces inventory carrying costs and frees up working capital "
            "tied up in safety stocks across the retail network."
        )
    elif "The $ values are illustrative" in p.text:
        p.text = (
            "These inventory buffer reductions highlight how model accuracy improvements directly lower safety stock "
            "requirements and improve capital efficiency for resource-constrained small businesses."
        )

# 3. Polish Introduction Sentences (P[7], P[8])
for p in doc.paragraphs:
    if "Cost and lack of in-house technical expertise are the top two reasons" in p.text:
        p.text = p.text.replace(
            "Cost and lack of in-house technical expertise are the top two reasons, not lack of awareness that tools are available, of why barriers to formal ERP and inventory-technology adoption exist, and this holds true in geographically diverse populations of SMEs, from micro-enterprises with limited access to affordable software [3] to city-level surveys where cost and management buy-in are top barriers to adoption [4].",
            "Cost and a lack of in-house technical expertise represent the primary barriers to formal ERP and inventory-technology adoption among SMEs, rather than a lack of awareness of available tools. This observation holds across geographically diverse small enterprises, ranging from micro-enterprises with limited software budgets [3] to urban SME surveys identifying cost and management buy-in as top adoption hurdles [4]."
        )
    if "There is some work in the area of micro and small enterprises now" in p.text:
        p.text = p.text.replace(
            "There is some work in the area of micro and small enterprises now that focuses directly on them [9] but there is a disconnect between the general purpose forecasting literature and a readily repeatable, low-cost forecasting system that is built around the constraints of micro and small enterprises.",
            "Recent studies have begun to focus directly on inventory forecasting for micro and small enterprises [9], but there remains a disconnect between general-purpose forecasting literature and a readily repeatable, low-cost forecasting system built around small business operational constraints."
        )

# Save to target file and copies
targets = [
    r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx',
    r'c:\Users\devan\Desktop\research paper\AI-Based Inventory Forecasting for Small Businesses.docx'
]

for t in targets:
    try:
        doc.save(t)
        print("Successfully saved final polished paper to:", t)
    except Exception as e:
        print(f"Error saving to {t}: {e}")
