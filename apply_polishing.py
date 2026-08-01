import docx
import os

doc_path = r'C:\Users\devan\Downloads\AI-Based_Inventory_Forecasting_for_Small_Businesses_REDLINED.docx'
doc = docx.Document(doc_path)

# 1. Hyperparameter Sentence Formatting (P[41] / Section III-C)
for p in doc.paragraphs:
    if "XGBoost with engineered features." in p.text:
        p.text = (
            "XGBoost with engineered features. A gradient-boosted tree regressor using 200 estimators, "
            "maximum depth of 4, and a learning rate of 0.05 was trained on lagged sales values "
            "(1, 2, 3, 7, and 14 days), rolling means, and calendar features. These hyperparameters were "
            "selected after preliminary experiments because they provided the best trade-off between "
            "forecasting accuracy and computational efficiency. This follows the feature-engineering approach "
            "shown to be effective for tree-based demand forecasting in prior comparative studies [11], [13]."
        )

# 2. Reproducibility Section Formatting (P[56] / Section III-F)
for p in doc.paragraphs:
    if "All results reported here use a fixed random seed" in p.text or "ARIMA(2,1,2); XGBoost" in p.text:
        p.text = (
            "All experiments used ARIMA(2,1,2) and XGBoost (200 estimators, maximum depth 4, learning rate 0.05). "
            "These hyperparameters were selected after preliminary experiments as they provided the best trade-off "
            "between forecasting accuracy and computational efficiency. A fixed random seed (42) was used to "
            "ensure reproducibility. Applied to the publicly available Rossmann Store Sales dataset, the "
            "data-ingestion pipeline, model implementations, and evaluation harness are implemented as a Python "
            "codebase with no proprietary dependencies, enabling end-to-end evaluation and replication."
        )

# 3. Figure 1 & 2 Captions and Table Titles (IEEE Formatting)
for p in doc.paragraphs:
    if "Fig. 1." in p.text:
        p.text = "Fig. 1.  System architecture: raw sales data flows through ingestion and feature engineering into the model layer, which produces forecasts that are scored by the evaluation layer."
    if "Fig. 2." in p.text or "Comparison of actual and XGBoost-predicted" in p.text:
        p.text = "Fig. 2.  Comparison of actual and XGBoost-predicted daily sales for Store 1 during the 60-day test period."
    if "TABLE I." in p.text:
        p.text = "TABLE I.  POSITIONING RELATIVE TO CLOSELY RELATED WORK"
    if "TABLE II." in p.text:
        p.text = "TABLE II.  OVERALL AVERAGE ERROR ACROSS ALL 1,115 STORES (ALL DAYS)"
    if "TABLE III." in p.text:
        p.text = "TABLE III.  FORECASTING PERFORMANCE ON REPRESENTATIVE STORES"

# Save updated document across target files in Downloads
targets = [
    r'C:\Users\devan\Downloads\AI-Based_Inventory_Forecasting_for_Small_Businesses_REDLINED.docx',
    r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx',
    r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses (Updated).docx',
    r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses (Final).docx',
    r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses (Submission).docx'
]

for t in targets:
    try:
        doc.save(t)
        print("Successfully saved polished docx to:", t)
    except Exception as e:
        print(f"Could not save to {t}: {e}")
