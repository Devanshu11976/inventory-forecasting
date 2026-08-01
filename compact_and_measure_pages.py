import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc_path = r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx'
doc = docx.Document(doc_path)

print(f"Loaded doc with {len(doc.paragraphs)} paragraphs.")

# 1. Remove all empty paragraphs except those containing images
paras_to_remove = []
for p in doc.paragraphs:
    if not p.text.strip() and 'blip' not in p._element.xml and 'graphic' not in p._element.xml:
        paras_to_remove.append(p)

for p in paras_to_remove:
    p._element.getparent().remove(p._element)

print(f"Removed {len(paras_to_remove)} empty paragraphs.")

# 2. Check if Figure 3 is already added; if not, insert Figure 3 after Table II / Section IV text
fig3_exists = any("Fig. 3." in p.text for p in doc.paragraphs)

if not fig3_exists:
    # Find paragraph containing TABLE II caption or Table II
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "TABLE II." in p.text or "OVERALL AVERAGE ERROR" in p.text:
            target_idx = i
            break
            
    if target_idx is not None:
        # Insert new paragraph after target_idx + offset
        insert_p = doc.paragraphs[target_idx + 1]
        
        # Add Figure 3 image paragraph
        fig3_p = insert_p.insert_paragraph_before()
        fig3_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fig3_p.add_run()
        img_path = r'c:\Users\devan\Desktop\research paper\figure3_model_comparison_ieee.png'
        run.add_picture(img_path, width=Inches(3.3)) # Fits single IEEE column or 3.3 in
        
        pf = fig3_p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        
        # Add Figure 3 caption paragraph
        cap_p = insert_p.insert_paragraph_before()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run("Fig. 3. Comparison of MAE, RMSE, MAPE, and sMAPE performance across baseline (MA-7), ARIMA(2,1,2), and proposed XGBoost models.")
        cap_run.font.name = 'Times New Roman'
        cap_run.font.size = Pt(8)
        
        cpf = cap_p.paragraph_format
        cpf.space_before = Pt(2)
        cpf.space_after = Pt(6)
        cpf.line_spacing = 1.0
        
        print("Successfully inserted Figure 3 chart and caption into document.")

# 3. Compact formatting across all paragraphs for 6-7 page target
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    has_img = 'blip' in p._element.xml or 'graphic' in p._element.xml
    
    if has_img:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        # Resize pictures inside runs if necessary
        for run in p.runs:
            for drawing in run._element.findall('.//' + qn('w:drawing')):
                # ensure drawing dimensions fit IEEE column width (~3.3 inches)
                extent = drawing.find('.//' + qn('wp:extent'))
                if extent is not None:
                    # 3.3 inches in EMUs = 3.3 * 914400 = 3017520
                    cx = int(extent.get('cx', '0'))
                    cy = int(extent.get('cy', '0'))
                    if cx > 3200000:
                        ratio = cy / cx
                        new_cx = 3017520
                        new_cy = int(new_cx * ratio)
                        extent.set('cx', str(new_cx))
                        extent.set('cy', str(new_cy))
    elif p.text.startswith('I.') or p.text.startswith('II.') or p.text.startswith('III.') or p.text.startswith('IV.') or p.text.startswith('V.') or p.text.startswith('VI.') or text == 'References':
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(3)
    elif text.startswith('A.') or text.startswith('B.') or text.startswith('C.') or text.startswith('D.'):
        pf = p.paragraph_format
        pf.space_before = Pt(5)
        pf.space_after = Pt(2)
    elif text.startswith('[') and text[1].isdigit():
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.0

# Save updated docx
out1 = r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx'
out2 = r'c:\Users\devan\Desktop\research paper\AI-Based Inventory Forecasting for Small Businesses.docx'

doc.save(out1)
doc.save(out2)
print("Saved compacted IEEE document.")
