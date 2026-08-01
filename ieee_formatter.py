#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
IEEE Conference Paper Formatter
Applies IEEE-standard formatting to the research paper .docx file.
Preserves all content; changes only visual formatting.
"""

import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import copy
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx'
doc = docx.Document(doc_path)
print(f"Loaded document: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.sections)} sections")

# ============================================================
# 1. MOVE SECTION BREAK FROM P[2] TO P[4]
#    Section 0 = Title / Author / Abstract / Index Terms (single-column)
#    Section 1 = Body text (two-column)
# ============================================================
p2_pPr = doc.paragraphs[2]._element.pPr
if p2_pPr is not None:
    sect_elem = p2_pPr.find(qn('w:sectPr'))
    if sect_elem is not None:
        saved_sect = copy.deepcopy(sect_elem)
        p2_pPr.remove(sect_elem)
        p4_pPr = doc.paragraphs[4]._element.get_or_add_pPr()
        p4_pPr.append(saved_sect)
        print("  Moved section break from P[2] → P[4]")

# ============================================================
# 2. PAGE SETUP — A4, IEEE margins
# ============================================================
A4_W = '11906'   # 210 mm in twips
A4_H = '16838'   # 297 mm in twips

for sec in doc.sections:
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(0.625)
    sec.right_margin  = Inches(0.625)
print("  Set A4 page size + IEEE margins on all sections")

# ============================================================
# 3. COLUMN LAYOUT
# ============================================================
# Section 0 (title block) — single column, continuous break
sect0_pr = doc.paragraphs[4]._element.pPr.find(qn('w:sectPr'))
if sect0_pr is not None:
    for tag in ['w:pgSz', 'w:pgMar', 'w:cols', 'w:type']:
        el = sect0_pr.find(qn(tag))
        if el is not None:
            sect0_pr.remove(el)
    sect0_pr.append(parse_xml(f'<w:type {nsdecls("w")} w:val="continuous"/>'))
    sect0_pr.append(parse_xml(f'<w:pgSz {nsdecls("w")} w:w="{A4_W}" w:h="{A4_H}"/>'))
    sect0_pr.append(parse_xml(
        f'<w:pgMar {nsdecls("w")} w:top="1080" w:bottom="1440" '
        f'w:left="900" w:right="900" w:header="720" w:footer="720"/>'))
    sect0_pr.append(parse_xml(f'<w:cols {nsdecls("w")} w:space="360"/>'))

# Section 1 (body) — two columns
sect1_pr = doc.sections[-1]._sectPr
cols_el = sect1_pr.find(qn('w:cols'))
if cols_el is not None:
    sect1_pr.remove(cols_el)
sect1_pr.append(parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>'))
print("  Set single-column title block + two-column body")

# ============================================================
# 4. HELPER — set font on one run
# ============================================================
def set_run_fmt(run, size=Pt(10), bold=None, italic=None, small_caps=None):
    run.font.name = 'Times New Roman'
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rF)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rF.set(qn(attr), 'Times New Roman')
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if small_caps is not None:
        sc = rPr.find(qn('w:smallCaps'))
        if small_caps:
            if sc is None:
                rPr.append(parse_xml(f'<w:smallCaps {nsdecls("w")}/>'))
        else:
            if sc is not None:
                rPr.remove(sc)
    # strip highlight / colour overrides
    for tag in ['w:color', 'w:highlight', 'w:shd']:
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)

# ============================================================
# 5. HELPER — format whole paragraph
# ============================================================
def fmt(p, size=Pt(10), bold=None, italic=None, small_caps=None,
        align=None, before=0, after=0,
        first_indent=None, hang=None, left=None,
        keep_next=False):
    for run in p.runs:
        set_run_fmt(run, size=size, bold=bold, italic=italic, small_caps=small_caps)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing  = 1.0
    # indentation
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if first_indent is not None or hang is not None or left is not None:
        if ind is None:
            ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
            pPr.append(ind)
        if first_indent is not None:
            ind.set(qn('w:firstLine'), str(first_indent))
            if qn('w:hanging') in ind.attrib:
                del ind.attrib[qn('w:hanging')]
        if hang is not None:
            ind.set(qn('w:hanging'), str(hang))
            if qn('w:firstLine') in ind.attrib:
                del ind.attrib[qn('w:firstLine')]
        if left is not None:
            ind.set(qn('w:left'), str(left))
    else:
        if ind is not None:
            for a in [qn('w:firstLine'), qn('w:hanging'), qn('w:left')]:
                if a in ind.attrib:
                    del ind.attrib[a]
    # keepNext (heading stays with following paragraph)
    if keep_next:
        if pPr.find(qn('w:keepNext')) is None:
            pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
    # widowControl on
    if pPr.find(qn('w:widowControl')) is None:
        pPr.append(parse_xml(f'<w:widowControl {nsdecls("w")}/>'))

# ============================================================
# 6. CLASSIFY + FORMAT EVERY PARAGRAPH
# ============================================================
RE_SEC  = re.compile(r'^[IVX]+\.\s+')
RE_SUB  = re.compile(r'^[A-Z]\.\s+')
RE_SSUB = re.compile(r'^\d+\)\s+')
RE_FIG  = re.compile(r'^Fig\.\s+\d+')
RE_TBL  = re.compile(r'^TABLE\s+[IVX]+')
RE_REF  = re.compile(r'^\[\d+\]')

prev_heading = False

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    has_img = 'blip' in p._element.xml

    # ---- empty paragraph (not an image) ----
    if not text and not has_img:
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        for r in p.runs:
            r.font.size = Pt(1)          # minimize empty-line height
        prev_heading = False
        continue

    # ---- image paragraph ----
    if has_img:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6); pf.space_after = Pt(3); pf.line_spacing = 1.0
        prev_heading = False
        continue

    # ---- title ----
    if i == 0:
        fmt(p, size=Pt(24), bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
        prev_heading = True; continue

    # ---- author / affiliation placeholders ----
    if i in [1, 2]:
        fmt(p, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER,
            before=0, after=0)
        prev_heading = False; continue

    # ---- abstract ----
    if text.startswith('Abstract'):
        runs = p.runs
        for r in runs:
            if 'Abstract' in r.text:
                set_run_fmt(r, size=Pt(9), bold=True, italic=False)
            else:
                set_run_fmt(r, size=Pt(9), bold=False, italic=False)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.0
        prev_heading = False; continue

    # ---- index terms ----
    if text.startswith('Index Terms'):
        runs = p.runs
        for r in runs:
            if 'Index Terms' in r.text:
                set_run_fmt(r, size=Pt(9), bold=True, italic=False)
            else:
                set_run_fmt(r, size=Pt(9), bold=False, italic=True)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(6); pf.line_spacing = 1.0
        prev_heading = False; continue

    # ---- section heading (I., II., …) or "References" ----
    if RE_SEC.match(text) or text == 'References':
        fmt(p, size=Pt(10), bold=True, small_caps=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            before=12, after=6, keep_next=True)
        prev_heading = True; continue

    # ---- subsection heading (A., B., …) ----
    if RE_SUB.match(text):
        fmt(p, size=Pt(10), bold=False, italic=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            before=9, after=3, keep_next=True)
        prev_heading = True; continue

    # ---- figure caption ----
    if RE_FIG.match(text):
        fmt(p, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER,
            before=3, after=6)
        prev_heading = False; continue

    # ---- table caption ----
    if RE_TBL.match(text):
        fmt(p, size=Pt(8), small_caps=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            before=6, after=3)
        prev_heading = False; continue

    # ---- reference entry ----
    if RE_REF.match(text):
        fmt(p, size=Pt(8),
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            before=0, after=1,
            left=360, hang=360)
        prev_heading = False; continue

    # ---- numbered list items 1) 2) 3) ----
    if RE_SSUB.match(text):
        fmt(p, size=Pt(10),
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            before=0, after=0,
            left=288, hang=288)
        prev_heading = False; continue

    # ---- body text ----
    if prev_heading:
        fmt(p, size=Pt(10), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            before=0, after=0, first_indent=0)
    else:
        fmt(p, size=Pt(10), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            before=0, after=0, first_indent=245)   # ~0.17 in

    prev_heading = False

print("  Formatted all paragraphs")

# ============================================================
# 7. FORMAT TABLES — 8pt, centered, bold header row
# ============================================================
for t_idx, table in enumerate(doc.tables):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        tblPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
    else:
        jc.set(qn('w:val'), 'center')

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_fmt(run, size=Pt(8),
                                bold=True if row_idx == 0 else None)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = para.paragraph_format
                pf.space_before = Pt(1); pf.space_after = Pt(1)
                pf.line_spacing = 1.0

print(f"  Formatted {len(doc.tables)} tables")

# ============================================================
# 8. SAVE
# ============================================================
out1 = r'C:\Users\devan\Downloads\AI-Based Inventory Forecasting for Small Businesses.docx'
out2 = r'c:\Users\devan\Desktop\research paper\AI-Based Inventory Forecasting for Small Businesses.docx'

for path in [out1, out2]:
    try:
        doc.save(path)
        print(f"  ✓ Saved → {path}")
    except Exception as e:
        print(f"  ✗ Could not save to {path}: {e}")

print("\nDone — IEEE formatting applied successfully.")
